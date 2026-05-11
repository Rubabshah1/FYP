#!/usr/bin/env python3
"""
ALKHIDMAT RAG SYSTEM - SUPABASE CLIENT EDITION
WITH DOMAIN CLASSIFICATION & CONFIDENCE SCORING
Uses Supabase Python client (REST API) and OpenAI GPT for all LLM calls.
"""

from dotenv import load_dotenv
load_dotenv()

import os
import re
import json
import time
import sys
from pathlib import Path
from typing import List, Dict, Tuple, Any, Optional
import uuid

import gc
import zipfile

import numpy as np
from scipy.stats import entropy
from sklearn.metrics.pairwise import cosine_similarity

# OpenAI
import openai

# Urdu helpers
from deep_translator import GoogleTranslator
import langdetect

# text splitter
from langchain_text_splitters import RecursiveCharacterTextSplitter

# Import domain anchors
from domain_anchors import DOMAIN_ANCHOR_QUERIES
from rag_config import (
    BRAND_TERMS,
    CHUNK_OVERLAP,
    CHUNK_SIZE,
    DOMAIN_CENTROID_REUSE_THRESHOLD,
    EMBEDDING_CACHE_ENABLE,
    EMBEDDING_CACHE_SIMILARITY_THRESHOLD,
    EMBEDDING_DIM,
    EMBEDDING_MODEL_NAME,
    EVIDENCE_COVERAGE_ENABLE,
    QUERY_ROUTER_ENABLE,
    RELEVANCE_THRESHOLD,
    RETRIEVAL_RETRY_ENABLE,
    RETRIEVAL_RETRY_MAX_ATTEMPTS,
    RETRIEVAL_RETRY_RELEVANCE_THRESHOLD,
    ROMAN_URDU_ENABLE,
    ROMAN_URDU_MARKERS,
    SELFRAG_ENABLE,
    SELFRAG_MIN_CONFIDENCE,
    SELFRAG_RELEVANCE_THRESHOLD,
    SELFRAG_RETRIEVE_THRESHOLD,
    SELFRAG_SUPPORT_THRESHOLD,
    SUPABASE_KEY,
    SUPABASE_URL,
    TRANSLATE_CONTEXT_FOR_URDU_OUTPUT,
)
from rag_embeddings import (
    cache_query_embedding,
    create_embeddings,
    find_similar_cached_embedding,
    get_embedder,
    normalize_query,
)
from rag_language import (
    build_query_lang_profile,
    analyze_query,
    select_retrieval_queries,
    is_urdu_script,
    looks_like_roman_urdu,
    translate_auto_to_english,
    translate_urdu_to_english,
    translate_english_to_urdu,
    protect_brand_terms,
    restore_brand_terms,
    to_roman_urdu,
)
from rag_supabase_client import get_supabase_client, test_connection

# ============================================================================
# GPT CONFIGURATION & UNIFIED CALLER
# ============================================================================

# Set OPENAI_API_KEY in your .env file
openai.api_key = os.environ.get("OPENAI_API_KEY", "")

# All LLM calls use this model. Override with GPT_MODEL env var.
GPT_MODEL = os.environ.get("GPT_MODEL", "gpt-4o-mini")


def _gpt_call(prompt: str, max_tokens: int = 50, temperature: float = 0.1,
              system: str = "You are a precise assistant. Follow instructions exactly.") -> str:
    """
    Unified GPT caller for ALL critic/classifier LLM calls in this file.
    Returns the response text string, or "" on error.
    """
    try:
        client = openai.OpenAI(api_key=openai.api_key)
        response = client.chat.completions.create(
            model=GPT_MODEL,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": prompt},
            ],
            max_tokens=max_tokens,
            temperature=temperature,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print(f"[GPT] Error calling {GPT_MODEL}: {e}", flush=True)
        return ""


def _gpt_generate(prompt: str, max_tokens: int = 400,
                  system: str = "You are a helpful assistant for Alkhidmat Foundation Pakistan.") -> str:
    """
    GPT caller for full answer generation (Step 4).
    Uses slightly higher temperature for natural answers.
    """
    try:
        client = openai.OpenAI(api_key=openai.api_key)
        response = client.chat.completions.create(
            model=GPT_MODEL,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": prompt},
            ],
            max_tokens=max_tokens,
            temperature=0.3,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print(f"[GPT] Generation error: {e}", flush=True)
        return ""


# ============================================================================
# SELF-RAG REFLECTION TOKENS & CRITIC
# ============================================================================

class SelfRAGReflectionTokens:
    RETRIEVE = "[Retrieve]"
    NO_RETRIEVE = "[No Retrieval]"
    RELEVANT = "[Relevant]"
    IRRELEVANT = "[Irrelevant]"
    FULLY_SUPPORTED = "[Fully supported]"
    PARTIALLY_SUPPORTED = "[Partially supported]"
    NO_SUPPORT = "[No support]"
    UTILITY_5 = "[Utility:5]"
    UTILITY_4 = "[Utility:4]"
    UTILITY_3 = "[Utility:3]"
    UTILITY_2 = "[Utility:2]"
    UTILITY_1 = "[Utility:1]"


class SelfRAGCritic:
    """All critic methods call GPT via _gpt_call. No local LLM used."""

    def is_domain_relevant(self, query: str) -> bool:
        query_lower = query.lower().strip()

        alkhidmat_keywords = [
            "alkhidmat", "alkhdmt", "alkhidmat foundation", "alkhidmat foundation pakistan"
        ]
        alkhidmat_programs = [
            "bano qabil", "banoqabil", "banu qabil",
            "aghosh", "aghosh home", "orphan",
            "mawakhat", "mawakhat program",
            "zakat", "sadaqah", "donation", "donate",
            "qurbani", "fidya", "kafara",
            "healthcare", "hospital", "clinic", "medical",
            "education", "scholarship",
            "disaster", "relief", "flood", "emergency",
            "volunteer", "volunteering"
        ]

        has_alkhidmat_keyword = any(kw in query_lower for kw in alkhidmat_keywords)
        has_program_keyword = any(p in query_lower for p in alkhidmat_programs)

        # Fast-path: if "Alkhidmat" appears in the query, it is ALWAYS relevant —
        # no GPT call needed. This covers contact/info queries like
        # "How can I contact Alkhidmat?" that GPT incorrectly marks as generic.
        if has_alkhidmat_keyword:
            print(f"[DOMAIN-RELEVANCE] Alkhidmat keyword found — skipping GPT, marking RELEVANT")
            return True

        if has_program_keyword:
            question_words = ["what", "who", "where", "when", "how", "why", "kya", "kaun", "kahan", "kab"]
            if any(query_lower.startswith(q) for q in question_words) or "?" in query:
                print(f"[DOMAIN-RELEVANCE] Program keyword in question: '{query}' — likely relevant")
            elif len(query_lower.split()) <= 5:
                print(f"[DOMAIN-RELEVANCE] Short query with program keyword: '{query}' — likely relevant")

        prompt = f"""Determine if this question is about Alkhidmat Foundation Pakistan.

Question: {query}

A query is RELEVANT if it refers to:
- Alkhidmat Foundation Pakistan by name, OR
- Alkhidmat's programs/services (Bano Qabil, Aghosh, Mawakhat, donations,
  healthcare, education, disaster relief, volunteering, etc.)
- Contacting, reaching, or getting information about Alkhidmat Foundation.
- Even if "Alkhidmat" is not explicitly mentioned, queries about the above programs are RELEVANT.

A query is IRRELEVANT if it is:
- About completely unrelated topics (politics, sports, celebrities, weather).
- General Islamic concepts with no connection to Alkhidmat.
- Too generic with no Alkhidmat context (e.g. "best charity in Pakistan", "free hospitals in Karachi").

When in doubt, choose [RELEVANT].

Respond with ONLY one of:
[RELEVANT]
[IRRELEVANT]

Answer:"""

        try:
            response = _gpt_call(prompt, max_tokens=10, temperature=0.1)
            if "[RELEVANT]" in response.upper():
                return True
            elif "[IRRELEVANT]" in response.upper():
                if has_program_keyword:
                    print(f"[DOMAIN-RELEVANCE] GPT said irrelevant but program keyword present — overriding to relevant")
                    return True
                return False
            return True  # default: relevant
        except Exception as e:
            print(f"[SelfRAG] Domain relevance check error: {e}")
            return True

    def should_retrieve(self, query: str) -> bool:
        prompt = f"""Determine if external knowledge retrieval is needed to answer this question.

Question: {query}

Consider:
- Does this require specific factual information about Alkhidmat Foundation?
- Is this about a specific organization, service, or policy?
- Can this be answered with general knowledge alone?

Respond with ONLY one of:
{SelfRAGReflectionTokens.RETRIEVE} - if retrieval is needed
{SelfRAGReflectionTokens.NO_RETRIEVE} - if general knowledge is sufficient

Answer:"""
        try:
            response = _gpt_call(prompt, max_tokens=10, temperature=0.1)
            if SelfRAGReflectionTokens.RETRIEVE in response:
                return True
            elif SelfRAGReflectionTokens.NO_RETRIEVE in response:
                return False
            return True
        except Exception as e:
            print(f"[SelfRAG] Retrieval prediction error: {e}")
            return True

    # NOTE: assess_relevance() (per-document LLM check) has been REMOVED.
    # Step 3 uses pure cosine-similarity threshold — faster and equally reliable.

    def check_answer_in_context(self, query: str, context: str) -> bool:
        """
        PRE-GENERATION GATE: Can the query be answered from this context at all?
        Prevents wasting GPT generation tokens when context is irrelevant.
        Different from verify_support — this runs BEFORE generation.
        """
        context_preview = context[:300] + "..." if len(context) > 1000 else context
        prompt = f"""Check if the provided context contains information to answer the question.

Question: {query}

Context:
{context_preview}

Does the context below contain information that is RELEVANT to answering the question?
Do not check if it answers perfectly — just check if it contains relevant facts.

Question: {query}
Context: {context_preview}

Respond ONLY:
[CAN_ANSWER] - if context contains relevant information
[CANNOT_ANSWER] - if context is completely unrelated

Answer:"""
        try:
            response = _gpt_call(prompt, max_tokens=10, temperature=0.1)
            if "[CAN_ANSWER]" in response.upper():
                return True
            elif "[CANNOT_ANSWER]" in response.upper():
                return False
            return True
        except Exception as e:
            print(f"[SelfRAG] Answer presence check error: {e}")
            return True

    def verify_support(self, query: str, answer: str, context: str) -> str:
        """
        POST-GENERATION HALLUCINATION CHECK: Is the generated answer grounded in context?

        WHY BOTH check_answer_in_context AND verify_support are kept:
        - check_answer_in_context (Step 3.5): runs BEFORE generation — gates whether
          we should generate at all ("is the answer findable in context?")
        - verify_support (Step 5): runs AFTER generation — checks whether the LLM
          actually stayed grounded ("did the generated answer make up facts?")
        They catch different failure modes and are both necessary.
        """
        prompt = f"""Verify if the answer is supported by the provided context.

Question: {query}

Context:
{context[:800]}

Answer:
{answer}

Check if the answer facts come from the context, or if the answer is making up information.

Respond with ONLY one of:
{SelfRAGReflectionTokens.FULLY_SUPPORTED}
{SelfRAGReflectionTokens.PARTIALLY_SUPPORTED}
{SelfRAGReflectionTokens.NO_SUPPORT}

Answer:"""
        try:
            response = _gpt_call(prompt, max_tokens=20, temperature=0.1)
            if SelfRAGReflectionTokens.FULLY_SUPPORTED in response:
                return "fully_supported"
            elif SelfRAGReflectionTokens.PARTIALLY_SUPPORTED in response:
                return "partially_supported"
            elif SelfRAGReflectionTokens.NO_SUPPORT in response:
                return "no_support"
            return "uncertain"
        except Exception as e:
            print(f"[SelfRAG] Support verification error: {e}")
            return "uncertain"

    def evaluate_utility(self, query: str, answer: str) -> int:
        """
        Utility judging always done in English to avoid mis-scoring Urdu answers.
        Translates Urdu/Roman-Urdu to English before calling GPT.
        """
        q_eval = query
        a_eval = answer
        if is_urdu_script(query) or is_urdu_script(answer):
            q_eval = translate_urdu_to_english(query) if is_urdu_script(query) else translate_auto_to_english(query)
            a_eval = translate_urdu_to_english(answer) if is_urdu_script(answer) else translate_auto_to_english(answer)

        prompt = f"""Evaluate how useful this answer is for the question.

Question: {q_eval}

Answer:
{a_eval}

Rate utility 1-5:
5 = Excellent, complete, directly answers the question
4 = Good, mostly answers the question
3 = Acceptable, provides some useful information
2 = Poor, barely addresses the question
1 = Very poor, does not answer the question

Respond with ONLY one of:
{SelfRAGReflectionTokens.UTILITY_5}
{SelfRAGReflectionTokens.UTILITY_4}
{SelfRAGReflectionTokens.UTILITY_3}
{SelfRAGReflectionTokens.UTILITY_2}
{SelfRAGReflectionTokens.UTILITY_1}

Answer:"""
        try:
            response = _gpt_call(prompt, max_tokens=15, temperature=0.1)
            if SelfRAGReflectionTokens.UTILITY_5 in response: return 5
            elif SelfRAGReflectionTokens.UTILITY_4 in response: return 4
            elif SelfRAGReflectionTokens.UTILITY_3 in response: return 3
            elif SelfRAGReflectionTokens.UTILITY_2 in response: return 2
            elif SelfRAGReflectionTokens.UTILITY_1 in response: return 1
            return 3
        except Exception as e:
            print(f"[SelfRAG] Utility evaluation error: {e}")
            return 3


# ============================================================================
# AGENTIC AI ARCHITECTURE
# ============================================================================



class RouterAgent:
    """Routes queries before embedding — decides if retrieval is needed. Uses GPT."""

    def __init__(self, critic: SelfRAGCritic):
        self.critic = critic

    def route(self, query: str) -> Tuple[bool, float, Optional[str]]:
        if not QUERY_ROUTER_ENABLE:
            return True, 1.0, None

        greetings = ["hello", "hi", "hey", "greetings", "salam", "assalam", "assalamu"]
        if any(query.lower().strip().startswith(g) for g in greetings):
            return False, 0.9, "Hello! How can I help you with information about Alkhidmat Foundation?"

        query_lower = query.lower().strip()
        question_words = ["what", "who", "where", "when", "how", "why", "which", "whose", "whom",
                          "kya", "kaun", "kahan", "kab", "kyun", "kis", "kisne"]
        question_patterns = ["are you", "is it", "do you", "can you", "will you", "would you",
                             "have you", "has it", "did you", "does it", "was it", "were you"]

        has_question_mark = "?" in query_lower
        starts_with_question = any(query_lower.startswith(q + " ") for q in question_words)
        has_question_pattern = any(p in query_lower for p in question_patterns)

        if not (has_question_mark or starts_with_question or has_question_pattern):
            nicety_patterns = [
                "thank you", "thanks", "thank", "thx", "ty", "shukriya", "shukria",
                "goodbye", "bye", "bye bye", "see you", "farewell", "khuda hafiz", "allah hafiz",
                "i appreciate", "appreciate it", "much appreciated", "grateful",
                "that's great", "that's good", "good job", "well done"
            ]
            if any(query_lower == p or query_lower.startswith(p + " ") for p in nicety_patterns):
                if any(query_lower.startswith(kw) for kw in
                       ["thank you", "thanks", "thank", "thx", "ty", "shukriya", "i appreciate", "appreciate"]):
                    return False, 0.9, "You're welcome! I'm glad I could help. Is there anything else you'd like to know about Alkhidmat Foundation?"
                elif any(query_lower.startswith(kw) for kw in ["bye", "goodbye", "see you", "farewell", "khuda hafiz"]):
                    return False, 0.9, "Goodbye! Feel free to come back if you have questions about Alkhidmat Foundation. Have a great day!"
                else:
                    return False, 0.9, "You're welcome! Is there anything else I can help you with?"

        retrieve_needed = self.critic.should_retrieve(query)  # GPT call
        if not retrieve_needed:
            return False, 0.7, None
        return True, 0.9, None


class RetrieverAgent:
    """Retrieves documents with retry logic and GPT-powered query reformulation."""

    def reformulate_query(self, original_query: str, domain: str, reason: str = "low_relevance") -> str:
        """GPT-powered query reformulation for low-relevance retry."""
        if not RETRIEVAL_RETRY_ENABLE:
            return original_query

        prompt = f"""Reformulate this query to improve document retrieval for Alkhidmat Foundation knowledge base. Reformulate this query as if a user is asking this query related to alkhidmat make it specific.

Original Query: {original_query}
Domain: {domain}
Reason: {reason}

Create a reformulated query that:
1. Includes key terms related to Alkhidmat Foundation
2. Adds domain-specific keywords (donation, healthcare, etc.)
3. Maintains the original intent
4. Is optimized for vector search

Respond with ONLY the reformulated query, nothing else."""

        try:
            reformulated = _gpt_call(prompt, max_tokens=60, temperature=0.3)
            print(f"[RETRIEVER-AGENT] Reformulated: '{original_query}' → '{reformulated}'")
            return reformulated if reformulated else original_query
        except Exception as e:
            print(f"[RETRIEVER-AGENT] Reformulation error: {e}")
            return original_query

    def reformulate_roman_urdu_query(self, roman_query: str, translated_en: str) -> str:
        """
        ALWAYS called for Roman Urdu queries.
        Auto-translation of Roman Urdu often produces awkward English.
        GPT rewrites it into clear, retrieval-optimised English.
        """
        prompt = f"""A user typed this query in Roman Urdu (Urdu written in English letters):
Roman Urdu original: {roman_query}
Auto-translated to English: {translated_en}

Rewrite the translated English into clear, natural English optimised for searching
an Alkhidmat Foundation Pakistan knowledge base. Keep the original intent intact.
Add relevant Alkhidmat-specific keywords if appropriate.

Respond with ONLY the rewritten English query, nothing else."""

        try:
            rewritten = _gpt_call(prompt, max_tokens=60, temperature=0.2)
            if rewritten:
                print(f"[ROMAN-UR] GPT-reformulated: '{roman_query}' → '{rewritten}'", flush=True)
                return rewritten
            return translated_en
        except Exception as e:
            print(f"[ROMAN-UR] Reformulation error: {e}")
            return translated_en

    def reformulate_for_alkhidmat_perspective(self, query_en: str) -> str:
        """
        ALWAYS called for every relevant query (after domain relevance check).
        Rewrites the query as if someone is asking Alkhidmat's own chatbot —
        even if they didn't mention "Alkhidmat" explicitly.
        Example: "what are your services?" → "What are Alkhidmat Foundation's services?"
        Example: "donate krna hai" → "How do I donate to Alkhidmat Foundation?"
        This dramatically improves retrieval for queries that omit the org name.
        """
        prompt = f"""You are helping rewrite a user query for Alkhidmat Foundation Pakistan's chatbot.

Original query: {query_en}

The user is talking TO Alkhidmat's chatbot. Rewrite their query so it explicitly 
mentions "Alkhidmat Foundation" where relevant, making it clear this is about 
Alkhidmat's specific services, programs, or operations.

Rules:
- If the query already mentions Alkhidmat explicitly, return it unchanged.
- Replace "you", "your", "you guys" with "Alkhidmat" / "Alkhidmat's".
- Add "Alkhidmat Foundation" context where it helps retrieval.
- Keep the original intent — do NOT add information.
- Keep it concise (under 20 words if possible).
- Return ONLY the rewritten query, nothing else.

Rewritten query:"""

        try:
            rewritten = _gpt_call(prompt, max_tokens=40, temperature=0.1)
            if rewritten and rewritten.strip():
                rewritten = rewritten.strip().strip("\"'")
                if rewritten.lower() != query_en.lower():
                    print(f"[PERSPECTIVE] Reformulated: '{query_en}' → '{rewritten}'", flush=True)
                return rewritten
            return query_en
        except Exception as e:
            print(f"[PERSPECTIVE] Reformulation error: {e}")
            return query_en

    def retrieve_with_retry(self, query: str, domain: str, top_k: int = 8,
                            filter_category: Optional[str] = None) -> Tuple[List[Dict], np.ndarray, List[np.ndarray], bool]:
        MIN_RELEVANT_DOCS = 2
        MIN_AVG_RELEVANCE = 0.6

        results, query_embedding, doc_embeddings = retrieve_from_supabase(
            query, top_k=top_k, filter_category=filter_category
        )
        was_retried = False

        if results:
            avg_relevance = sum(r.get('similarity', 0) for r in results) / len(results)
            high_relevance_count = sum(1 for r in results if r.get('similarity', 0) >= MIN_AVG_RELEVANCE)

            if high_relevance_count >= MIN_RELEVANT_DOCS and avg_relevance >= MIN_AVG_RELEVANCE:
                if not RETRIEVAL_RETRY_ENABLE:
                    return results, query_embedding, doc_embeddings, was_retried
                if avg_relevance >= RETRIEVAL_RETRY_RELEVANCE_THRESHOLD:
                    return results, query_embedding, doc_embeddings, was_retried

                print(f"[RETRIEVER-AGENT] ⚠️ Low relevance ({avg_relevance:.3f}), attempting GPT reformulation...")
                for attempt in range(RETRIEVAL_RETRY_MAX_ATTEMPTS):
                    reformulated = self.reformulate_query(query, domain, reason="low_relevance")
                    cached_emb = find_similar_cached_embedding(reformulated)
                    if cached_emb is not None:
                        query_embedding = cached_emb
                    else:
                        embedder = get_embedder()
                        query_embedding = embedder.encode([f"query: {reformulated}"], normalize_embeddings=True)[0]
                        cache_query_embedding(reformulated, query_embedding)

                    retry_results, _, retry_doc_embeddings = retrieve_from_supabase(
                        reformulated, top_k=top_k, filter_category=filter_category, query_embedding=query_embedding
                    )
                    if retry_results:
                        retry_avg = sum(r.get('similarity', 0) for r in retry_results) / len(retry_results)
                        if retry_avg > avg_relevance:
                            print(f"[RETRIEVER-AGENT] ✅ Reformulation improved: {avg_relevance:.3f} → {retry_avg:.3f}")
                            was_retried = True
                            return retry_results, query_embedding, retry_doc_embeddings, was_retried
                        else:
                            print(f"[RETRIEVER-AGENT] ⚠️ Reformulation did not improve relevance")
                return results, query_embedding, doc_embeddings, was_retried

            is_domain_specific = (filter_category and filter_category != "general") or \
                                  (not filter_category and domain != "general")
            if is_domain_specific:
                print(f"[RETRIEVER-AGENT] ⚠️ Only {high_relevance_count} relevant doc(s) in {domain} — falling back to general...")
                general_results, _, general_doc_embeddings = retrieve_from_supabase(
                    query, top_k=top_k, filter_category="general", query_embedding=query_embedding
                )
                if general_results:
                    seen = set()
                    combined = []
                    for r in results + general_results:
                        key = (r.get('file_path', ''), r.get('chunk_index', 0))
                        if key not in seen:
                            seen.add(key)
                            combined.append(r)
                    combined = sorted(combined, key=lambda x: x.get('similarity', 0), reverse=True)[:top_k]
                    combined_embeddings = doc_embeddings.copy()
                    combined_embeddings.extend(general_doc_embeddings)
                    print(f"[RETRIEVER-AGENT] ✅ Combined {len(results)} {domain} + {len(general_results)} general → {len(combined)}")
                    was_retried = True
                    return combined, query_embedding, combined_embeddings, was_retried

        if not results:
            is_domain_specific = (filter_category and filter_category != "general") or \
                                  (not filter_category and domain != "general")
            if is_domain_specific:
                print(f"[RETRIEVER-AGENT] ⚠️ No results in {domain} — falling back to general...")
                general_results, _, general_doc_embeddings = retrieve_from_supabase(
                    query, top_k=top_k, filter_category="general", query_embedding=query_embedding
                )
                if general_results:
                    print(f"[RETRIEVER-AGENT] ✅ Found {len(general_results)} in general domain")
                    was_retried = True
                    return general_results, query_embedding, general_doc_embeddings, was_retried

        return results, query_embedding, doc_embeddings, was_retried


class EvidenceCoverageAgent:
    """Claim-by-claim hallucination check using GPT."""

    def __init__(self, critic: SelfRAGCritic):
        self.critic = critic

    def _split_compound_claims(self, sentence: str) -> List[str]:
        sentence = sentence.strip()
        if not sentence or len(sentence) < 10:
            return []

        subject_match = re.match(
            r'^([A-Z][^,\.!?]+?)(?:\s+(?:is|are|was|were|provides|provide|works|work|supports|support|offers|offer|operates|operate))',
            sentence)
        subject = subject_match.group(1).strip() if subject_match else None

        if re.search(r',\s+(?:and|or)\s+', sentence, re.IGNORECASE):
            parts = re.split(r',\s+(?:and|or)\s+', sentence, flags=re.IGNORECASE)
            if len(parts) > 1:
                claims = []
                for i, part in enumerate(parts):
                    part = re.sub(r'^,\s*', '', part.strip())
                    if i > 0 and subject and part and not part[0].isupper():
                        if not re.search(r'\b(?:provides|provide|works|work|supports|support|offers|offer|operates|operate|is|are)\b', part, re.IGNORECASE):
                            part = f"{subject} {part}"
                    if len(part.strip()) > 10:
                        claims.append(part.strip())
                if claims:
                    return claims

        if re.search(r'\s+and\s+', sentence, re.IGNORECASE) and not re.search(r',\s+and\s+', sentence, re.IGNORECASE):
            parts = re.split(r'\s+and\s+', sentence, flags=re.IGNORECASE)
            if len(parts) > 1:
                claims = []
                for i, part in enumerate(parts):
                    part = part.strip()
                    if i > 0 and subject and part and not part[0].isupper():
                        if not re.search(r'\b(?:provides|provide|works|work|supports|support|offers|offer|operates|operate|is|are)\b', part, re.IGNORECASE):
                            part = f"{subject} {part}"
                    if len(part.strip()) > 10:
                        claims.append(part.strip())
                if claims:
                    return claims

        comma_parts = [p.strip() for p in sentence.split(',')]
        if len(comma_parts) > 2 and sum(len(p) for p in comma_parts) / len(comma_parts) < 50:
            claims = [comma_parts[0]] if comma_parts[0] else []
            for part in comma_parts[1:]:
                part = re.sub(r'\s+(?:and|or)\s*$', '', part.strip(), flags=re.IGNORECASE)
                if subject and part and not part[0].isupper():
                    if not re.search(r'\b(?:provides|provide|works|work|supports|support|offers|offer|operates|operate|is|are)\b', part, re.IGNORECASE):
                        part = f"{subject} {part}"
                if len(part.strip()) > 10:
                    claims.append(part.strip())
            if len(claims) > 1:
                return claims

        relative_match = re.search(r'(.+?)\s+(?:that|which|who)\s+(.+)', sentence)
        if relative_match:
            main_clause = relative_match.group(1).strip()
            relative_clause = relative_match.group(2).strip()
            claims = []
            if len(main_clause) > 10:
                claims.append(main_clause)
            if len(relative_clause) > 10:
                if subject and not relative_clause[0].isupper():
                    relative_clause = f"{subject} {relative_clause}"
                claims.append(relative_clause)
            if claims:
                return claims

        return [sentence] if len(sentence.strip()) > 10 else []

    def check_coverage(self, answer: str, context: str, query: str, domain: str = "general",
                       results: Optional[List[Dict]] = None) -> Tuple[bool, List[str], float]:
        if not EVIDENCE_COVERAGE_ENABLE:
            return True, [], 1.0

        # Skip coverage check for "I don't know" / fallback responses — nothing to verify
        _fallback_phrases = [
            "i don't know", "i do not know", "مجھے معلوم نہیں",
            "i don't have that information", "i cannot provide",
            "i apologize", "no information", "not found"
        ]
        if any(phrase in answer.lower() for phrase in _fallback_phrases):
            print(f"[EVIDENCE-COVERAGE] Skipping — answer is a fallback/I-don't-know response ✅")
            return True, [], 1.0

        allow_abstractive = self._is_summary_allowed(query, domain)
        if allow_abstractive:
            print(f"[EVIDENCE-COVERAGE] Summary mode — relaxing strictness for {domain} domain")

        sentences = [s.strip() for s in re.split(r'[.!?]\s+', answer) if len(s.strip()) > 10]
        if not sentences:
            return True, [], 1.0

        all_claims = []
        for sentence in sentences:
            all_claims.extend(self._split_compound_claims(sentence))
        if not all_claims:
            return True, [], 1.0

        print(f"[EVIDENCE-COVERAGE] {len(all_claims)} atomic claims from {len(sentences)} sentences")

        document_contexts = None
        if results and len(results) > 1:
            document_contexts = [r.get('text', '') for r in results]
            print(f"[EVIDENCE-COVERAGE] Multi-document support ({len(document_contexts)} docs)")

        unsupported = []
        supported_count = 0

        for claim in all_claims:
            is_supported, confidence = self._check_claim_support(
                claim, context, query,
                document_contexts=document_contexts,
                allow_abstractive=allow_abstractive
            )
            threshold = 0.3 if allow_abstractive else 0.4
            if not is_supported or confidence < threshold:
                unsupported.append(claim)
            else:
                supported_count += 1

        coverage_score = supported_count / len(all_claims) if all_claims else 1.0
        all_covered = coverage_score >= (0.5 if allow_abstractive else 0.4)

        if not all_covered:
            print(f"[EVIDENCE-COVERAGE] ⚠️ {len(unsupported)}/{len(all_claims)} unsupported (coverage: {coverage_score:.2f})")
            for claim in unsupported[:3]:
                print(f"   - '{claim[:80]}'")
        else:
            print(f"[EVIDENCE-COVERAGE] ✅ All supported (coverage: {coverage_score:.2f})")

        return all_covered, unsupported, coverage_score

    def _is_summary_allowed(self, query: str, domain: str) -> bool:
        summary_patterns = ["what is", "who is", "tell me about", "describe", "about",
                            "overview", "introduction", "summary", "kya hai", "kya hain", "kya hota hai"]
        return (any(p in query.lower() for p in summary_patterns) or
                domain.lower() in ["general", "about", "introduction"])

    def _check_claim_support(self, claim: str, context: str, query: str,
                             document_contexts: Optional[List[str]] = None,
                             allow_abstractive: bool = False) -> Tuple[bool, float]:
        if document_contexts and len(document_contexts) > 1:
            full_context = "\n\n".join(
                [f"Document {i+1}:\n{dc[:500]}" for i, dc in enumerate(document_contexts[:5])]
            )
        else:
            full_context = context[:2000]

        leniency = ("Abstractive synthesis is acceptable. Be lenient — if context is topically related, choose [SUPPORTED]."
                    if allow_abstractive else
                    "Be lenient — if context is topically related to the claim, choose [SUPPORTED].")

        prompt = f"""Check if this claim is supported by the provided context.

Claim: {claim}

Context:
{full_context}

{leniency}

Answer with ONLY:
[SUPPORTED]
[NOT_SUPPORTED]

Answer:"""
        try:
            response = _gpt_call(prompt, max_tokens=15, temperature=0.1)
            if "[SUPPORTED]" in response.upper():
                return True, 0.8
            elif "[NOT_SUPPORTED]" in response.upper():
                claim_kw = set(claim.lower().split())
                ctx_kw = set(full_context.lower().split())
                overlap = len(claim_kw & ctx_kw) / len(claim_kw) if claim_kw else 0
                return (True, 0.4) if overlap > 0.3 else (False, 0.2)
            return True, 0.6
        except Exception as e:
            print(f"[EVIDENCE-COVERAGE] Error: {e}")
            return True, 0.5




# ============================================================================
# DOMAIN CLASSIFICATION CLASS
# ============================================================================

class DomainClassifier:
    """Embedding-based domain classification. No LLM call needed."""

    _domain_embeddings_cache = None
    _embedding_model = None

    @staticmethod
    def initialize_domain_embeddings(model_name: str = None):
        if DomainClassifier._domain_embeddings_cache is not None:
            return
        if os.environ.get('BATCH_MODE') != 'True':
            print("\n🔄 Initializing domain embeddings...")

        model = get_embedder()
        DomainClassifier._embedding_model = model
        domain_embeddings = {}

        for domain, queries in DOMAIN_ANCHOR_QUERIES.items():
            prefixed = [f"query: {q}" for q in queries]
            embeddings = model.encode(prefixed, show_progress_bar=False, normalize_embeddings=True)
            centroid = np.mean(embeddings, axis=0)
            domain_embeddings[domain] = centroid
            if os.environ.get('BATCH_MODE') != 'True':
                print(f" ✓ {domain}: {len(queries)} anchors → centroid computed")

        DomainClassifier._domain_embeddings_cache = domain_embeddings
        if os.environ.get('BATCH_MODE') != 'True':
            print("✅ Domain embeddings initialized!\n")

    @staticmethod
    def classify_domain(query: str) -> Dict[str, any]:
        if DomainClassifier._domain_embeddings_cache is None:
            DomainClassifier.initialize_domain_embeddings()

        query_embedding = DomainClassifier._embedding_model.encode(
            [f"query: {query}"], normalize_embeddings=True
        )[0]
        similarities = {
            domain: float(cosine_similarity(query_embedding.reshape(1, -1), centroid.reshape(1, -1))[0][0])
            for domain, centroid in DomainClassifier._domain_embeddings_cache.items()
        }
        winning_domain = max(similarities, key=similarities.get)
        return {'domain': winning_domain, 'confidence': similarities[winning_domain], 'all_scores': similarities}

    @staticmethod
    def get_domain_emoji(domain: str) -> str:
        return {'donation': '💰', 'healthcare': '🏥', 'general': '📋'}.get(domain, '📋')


# ============================================================================
# CONFIDENCE SCORING
# ============================================================================

class ConfidenceScorer:
    """
    Confidence scoring based on retrieval similarity.
    Note: token log-probs are NOT available from the GPT API.
    Retrieval confidence (cosine similarity of query vs docs) is the primary signal.
    """

    @staticmethod
    def calculate_retrieval_confidence(query_embedding: np.ndarray,
                                       doc_embeddings: List[np.ndarray],
                                       top_k: int = 8) -> float:
        if not doc_embeddings:
            return 0.0
        sims = [float(cosine_similarity(query_embedding.reshape(1, -1), doc_emb.reshape(1, -1))[0][0])
                for doc_emb in doc_embeddings[:top_k]]
        return float(np.mean(sims))

    @staticmethod
    def calculate_combined_confidence(retrieval_confidence: float,
                                      selfrag_scores: Dict = None,
                                      log_probs: List[float] = None,
                                      token_probs_distributions: List[np.ndarray] = None) -> Dict[str, float]:
        scores = {'retrieval_confidence': retrieval_confidence}
        if selfrag_scores and selfrag_scores.get('relevance_score', 0) > 0:
            scores['selfrag_relevance'] = selfrag_scores['relevance_score']
        scores['combined_confidence'] = retrieval_confidence
        return scores


def _print_confidence_scores(confidence_scores: Dict, label: str = "CONFIDENCE SCORES"):
    """Print confidence scores in a consistent formatted block."""
    print(f"\n{'='*80}", flush=True)
    print(f"{label}:", flush=True)
    print(f"{'='*80}", flush=True)
    combined = confidence_scores.get('combined_confidence', 0)
    retrieval = confidence_scores.get('retrieval_confidence', 0)
    selfrag_rel = confidence_scores.get('selfrag_relevance', None)

    print(f" Combined Confidence   : {combined:.4f} ⭐", flush=True)
    print(f" ├─ Retrieval Conf.    : {retrieval:.4f}  (query↔docs cosine similarity)", flush=True)
    if selfrag_rel is not None:
        print(f" └─ SelfRAG Relevance  : {selfrag_rel:.4f}", flush=True)
    print(f" [Note: token log-probs not available with GPT API]", flush=True)

    if combined >= 0.7:
        print(f"✅ High confidence — answer is likely reliable", flush=True)
    elif combined >= 0.5:
        print(f"⚠️  Moderate confidence — answer may need verification", flush=True)
    else:
        print(f"❌ Low confidence — answer should be verified from sources", flush=True)
    print(f"{'='*80}\n", flush=True)
    sys.stdout.flush()


# ============ Language helpers (local, no LLM) ============

def protect_brand_terms(text: str) -> str:
    for term in BRAND_TERMS:
        text = re.sub(rf"\b{re.escape(term)}\b", f"@@{term}@@", text, flags=re.IGNORECASE)
    return text

def restore_brand_terms(text: str) -> str:
    return text.replace("@@", "")

def is_urdu_script(text: str) -> bool:
    return bool(re.search(r'[\u0600-\u06FF]', text))

def looks_like_roman_urdu(text: str) -> bool:
    if is_urdu_script(text) or re.search(r'[^\x00-\x7F]', text):
        return False
    tokens = re.findall(r"[a-zA-Z']+", text.lower())
    if not tokens:
        return False
    hits = sum(1 for t in tokens if t in ROMAN_URDU_MARKERS)
    return hits >= 2 or (hits >= 1 and len(tokens) <= 6)

def translate_auto_to_english(text: str) -> str:
    try:
        return GoogleTranslator(source='auto', target='en').translate(text)
    except Exception:
        return text

def detect_language(text: str) -> str:
    try:
        return langdetect.detect(text)
    except Exception:
        return "en"

def is_urdu(text: str) -> bool:
    return bool(re.compile(r'[\u0600-\u06FF]').search(text)) or detect_language(text) == "ur"

def translate_urdu_to_english(text: str) -> str:
    try:
        return GoogleTranslator(source='ur', target='en').translate(text)
    except Exception:
        return text

def translate_english_to_urdu(text: str, timeout: int = 10) -> str:
    try:
        from concurrent.futures import ThreadPoolExecutor, TimeoutError as FutureTimeoutError
        with ThreadPoolExecutor(max_workers=1) as executor:
            future = executor.submit(lambda: GoogleTranslator(source='en', target='ur').translate(text))
            try:
                result = future.result(timeout=timeout)
                return result if result else text
            except FutureTimeoutError:
                print(f"[WARNING] Translation timed out after {timeout}s.", flush=True)
                return text
    except Exception as e:
        print(f"[WARNING] Translation failed: {e}", flush=True)
        return text


class QueryLangProfile:
    def __init__(self, original_query: str, input_lang: str, query_en: str,
                 output_lang: str, query_urdu_script: Optional[str] = None):
        self.original_query = original_query
        self.input_lang = input_lang
        self.query_en = query_en
        self.output_lang = output_lang
        self.query_urdu_script = query_urdu_script


def normalize_implicit_org_references(query: str) -> str:
    """Replace 'you/your/you guys' → 'Alkhidmat/Alkhidmat's' for better retrieval."""
    q = query.strip()
    for pattern, replacement in [
        (r"\byou guys\b", "Alkhidmat"),
        (r"\byour\b", "Alkhidmat's"),
        (r"\byou\b", "Alkhidmat"),
    ]:
        q = re.sub(pattern, replacement, q, flags=re.IGNORECASE)
    return q


def build_query_lang_profile(query: str) -> QueryLangProfile:
    """
    Build language profile.
    Roman Urdu queries are ALWAYS GPT-reformulated for better retrieval quality.
    """
    q = query.strip()

    if is_urdu_script(q) or detect_language(q) == "ur":
        q_en = translate_urdu_to_english(q)
        q_en = normalize_implicit_org_references(q_en)
        return QueryLangProfile(original_query=q, input_lang="ur", query_en=q_en,
                                output_lang="ur", query_urdu_script=q)

    if looks_like_roman_urdu(q):
        # Step 1: auto-translate
        q_en_raw = translate_auto_to_english(q)
        q_en_raw = normalize_implicit_org_references(q_en_raw)
        # Step 2: ALWAYS GPT-reformulate Roman Urdu — auto-translation is often awkward
        retriever_helper = RetrieverAgent()
        q_en = retriever_helper.reformulate_roman_urdu_query(q, q_en_raw)
        return QueryLangProfile(original_query=q, input_lang="roman_ur", query_en=q_en,
                                output_lang="roman_ur", query_urdu_script=None)

    q_en = normalize_implicit_org_references(q)
    return QueryLangProfile(original_query=q, input_lang="en", query_en=q_en,
                            output_lang="en", query_urdu_script=None)


# ============ Query Analysis ============
def analyze_query(query: str) -> Dict[str, Any]:
    q_lower = query.lower().strip()
    procedural = ["how to", "how do i", "how can i", "steps", "step by step", "procedure"]
    roman = ["kaise", "kesy", "kese", "kaisay", "kya", "kyun", "kyu"]
    list_kw = ['list', 'points', 'bullet', 'enumerate', 'ways', 'methods']
    summary_kw = ['summarize', 'summary', 'briefly', 'overview', 'خلاصہ', 'khulasa']
    detail_kw = ['explain', 'detail', 'describe', 'why', 'تفصیل', 'tafseel']

    wants_steps = any(m in q_lower for m in procedural) or any(m in q_lower for m in roman)
    return {
        'wants_list': wants_steps or any(kw in q_lower for kw in list_kw),
        'wants_summary': any(kw in q_lower for kw in summary_kw),
        'wants_detail': any(kw in q_lower for kw in detail_kw),
        'is_urdu': is_urdu(query)
    }

def expand_query_for_retrieval(query_en: str, domain: str, query_info: Dict[str, Any]) -> str:
    q = query_en.strip()
    if len(q.split()) >= 6 and not query_info.get("wants_list"):
        return q
    if domain == "donation":
        return f"{q}. donate donation methods how to donate steps JazzCash EasyPaisa bank transfer online donation international account"
    if domain == "healthcare":
        return f"{q}. Alkhidmat hospital clinic services eligibility locations how to get treatment"
    return q


# ============ Document Processing ============

def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        import PyPDF2
        from io import BytesIO
        pdf_reader = PyPDF2.PdfReader(BytesIO(file_bytes))
        parts = []
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                text = page.extract_text()
                if text.strip():
                    parts.append(text)
            except Exception as e:
                print(f"   ⚠️ Page {page_num+1} error: {e}")
        return "\n\n".join(parts)
    except ImportError:
        print("   ⚠️ PyPDF2 not installed.")
        return ""
    except Exception as e:
        print(f"   ⚠️ PDF error: {e}")
        return ""

def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
        from io import BytesIO
        doc = Document(BytesIO(file_bytes))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = [c.text.strip() for c in row.cells if c.text.strip()]
                if row_text:
                    parts.append(" | ".join(row_text))
        return "\n\n".join(parts)
    except ImportError:
        print("   ⚠️ python-docx not installed.")
        return ""
    except Exception as e:
        print(f"   ⚠️ DOCX error: {e}")
        return ""

def load_documents_from_zip(zip_path: str, file_paths_filter: Optional[set] = None) -> Dict[str, List[Dict]]:
    if not os.path.exists(zip_path):
        raise FileNotFoundError(f"ZIP not found: {zip_path}")
    documents_by_category = {}
    skipped = []

    with zipfile.ZipFile(zip_path, "r") as zip_ref:
        for file_path in zip_ref.namelist():
            if file_path.endswith("/"):
                continue
            if "__MACOSX" in file_path or Path(file_path).name.startswith("._"):
                skipped.append(file_path)
                continue
            if file_paths_filter is not None and file_path not in file_paths_filter:
                continue
            fname = Path(file_path).name.lower()
            if fname in [".ds_store", "thumbs.db", ".gitignore", ".gitkeep"]:
                skipped.append(file_path)
                continue
            ext = Path(file_path).suffix.lower()
            if ext not in [".txt", ".pdf", ".docx"]:
                skipped.append(file_path)
                continue
            parts = Path(file_path).parts
            if len(parts) < 3:
                skipped.append(file_path)
                continue
            category = parts[-2]
            filename = parts[-1]
            try:
                with zip_ref.open(file_path) as f:
                    file_bytes = f.read()
                if ext == ".txt":
                    try:
                        content = file_bytes.decode("utf-8")
                    except UnicodeDecodeError:
                        content = file_bytes.decode("latin-1")
                elif ext == ".pdf":
                    content = extract_text_from_pdf(file_bytes)
                    if not content.strip():
                        skipped.append(file_path)
                        continue
                elif ext == ".docx":
                    content = extract_text_from_docx(file_bytes)
                    if not content.strip():
                        skipped.append(file_path)
                        continue
                else:
                    skipped.append(file_path)
                    continue
                if content.strip():
                    documents_by_category.setdefault(category, []).append({
                        "content": content, "filename": filename,
                        "category": category, "file_path": file_path
                    })
                else:
                    skipped.append(file_path)
            except Exception as e:
                print(f"   ⚠️ Error processing {file_path}: {e}")
                skipped.append(file_path)

    if skipped:
        print(f"⚠️ Skipped {len(skipped)} files")

    total = sum(len(docs) for docs in documents_by_category.values())
    print(f"✅ Loaded {total} documents from ZIP")
    return documents_by_category

def clean_text(text: str) -> str:
    text = re.sub(r'={3,}[\s\S]*?={3,}', '', text)
    text = re.sub(r'URL:\s*https?://\S+', '', text)
    text = re.sub(r'TITLE:.*?\n', '', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    return text.strip()

def prepare_documents(zip_path: str) -> Tuple[List[str], List[Dict]]:
    docs_by_cat = load_documents_from_zip(zip_path)
    all_docs, metadata = [], []
    for cat, docs in docs_by_cat.items():
        for doc in docs:
            c = clean_text(doc["content"])
            if c:
                all_docs.append(c)
                metadata.append({"filename": doc["filename"], "category": doc["category"],
                                  "file_path": doc["file_path"]})
    print(f"Prepared {len(all_docs)} documents")
    return all_docs, metadata

def split_documents(documents: List[str], metadata: List[Dict]):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP, length_function=len,
        separators=["\n\n\n", "\n\n", "\n", ". ", "! ", "? ", "؟ ", "۔ ", " ", ""],
    )
    chunks, metas = [], []
    for doc, meta in zip(documents, metadata):
        for idx, p in enumerate(splitter.split_text(doc)):
            chunks.append(p)
            chunk_meta = meta.copy()
            chunk_meta["chunk_index"] = idx
            metas.append(chunk_meta)
    print(f"Split into {len(chunks)} chunks (avg {int(np.mean([len(c) for c in chunks]))} chars)")
    return chunks, metas


# ============ Supabase Storage ============

def save_chunks_to_supabase(chunks: List[str], metadata: List[Dict], embeddings: np.ndarray):
    supabase = get_supabase_client()
    print("Saving to Supabase...")
    rows = [{"doc_id": str(uuid.uuid4()), "chunk_text": chunk,
             "chunk_index": metadata[i].get("chunk_index", 0), "category": metadata[i].get("category"),
             "filename": metadata[i].get("filename"), "file_path": metadata[i].get("file_path"),
             "doc_domain": metadata[i].get("category"), "embedding": embeddings[i].tolist()}
            for i, chunk in enumerate(chunks)]
    batch_size = 100
    total_inserted = 0
    for i in range(0, len(rows), batch_size):
        batch = rows[i:i + batch_size]
        try:
            supabase.table("documents").insert(batch).execute()
            total_inserted += len(batch)
            print(f" Batch {i//batch_size+1}: {len(batch)} chunks")
        except Exception as e:
            print(f" ⚠️ Batch error: {e}")
            for row in batch:
                try:
                    supabase.table("documents").insert(row).execute()
                    total_inserted += 1
                except Exception as e2:
                    print(f" ⚠️ Row error: {e2}")
    print(f"✅ Stored {total_inserted}/{len(rows)} chunks")

def clear_documents_table():
    supabase = get_supabase_client()
    try:
        supabase.table("documents").delete().neq("doc_id", "00000000-0000-0000-0000-000000000000").execute()
        print("✅ Cleared all documents")
    except Exception as e:
        print(f"⚠️ Error clearing: {e}")

def document_exists(file_path: str) -> bool:
    try:
        result = get_supabase_client().table("documents").select("doc_id").eq("file_path", file_path).limit(1).execute()
        return len(result.data) > 0 if result.data else False
    except Exception as e:
        print(f"⚠️ Existence check error: {e}")
        return False

def delete_document_by_path(file_path: str) -> bool:
    try:
        result = get_supabase_client().table("documents").delete().eq("file_path", file_path).execute()
        print(f"✅ Deleted {len(result.data) if result.data else 0} chunks for: {file_path}")
        return True
    except Exception as e:
        print(f"⚠️ Delete error: {e}")
        return False

def list_knowledge_base_documents():
    """List all unique documents in the knowledge base."""
    supabase = get_supabase_client()
    try:
        # Fetch only necessary columns to identify unique documents
        res = supabase.table("documents").select("file_path, filename, category").execute()
        if not res.data:
            return []
        
        docs = {}
        for row in res.data:
            path = row['file_path']
            if path not in docs:
                docs[path] = {
                    "file_path": path,
                    "filename": row['filename'],
                    "category": row['category'],
                    "chunks": 0
                }
            docs[path]["chunks"] += 1
            
        # Convert dict to sorted list by filename
        result = list(docs.values())
        result.sort(key=lambda x: x['filename'].lower())
        return result
    except Exception as e:
        print(f"⚠️ List error: {e}")
        return []

def add_document_incremental(file_path: str, content: str, category: str, filename: str,
                             reindex: bool = False) -> bool:
    supabase = get_supabase_client()
    if document_exists(file_path):
        if reindex:
            print(f"🔄 Re-indexing: {file_path}")
            delete_document_by_path(file_path)
        else:
            print(f"⏭️ Already exists, skipping: {file_path}")
            return True

    print(f"\n📄 Processing: {filename} [{category}]")
    cleaned = clean_text(content)
    if not cleaned:
        print(f"⚠️ No content after cleaning, skipping")
        return False

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP, length_function=len,
        separators=["\n\n\n", "\n\n", "\n", ". ", "! ", "? ", "؟ ", "۔ ", " ", ""]
    )
    chunks = splitter.split_text(cleaned)
    print(f"   Split into {len(chunks)} chunks")
    if not chunks:
        return False

    embedder = get_embedder()
    embeddings = np.array(embedder.encode([f"passage: {c}" for c in chunks],
                                          show_progress_bar=False, batch_size=32,
                                          normalize_embeddings=True)).astype("float32")
    rows = [{"doc_id": str(uuid.uuid4()), "chunk_text": chunk, "chunk_index": idx,
             "category": category, "filename": filename, "file_path": file_path,
             "doc_domain": category, "embedding": embeddings[idx].tolist()}
            for idx, chunk in enumerate(chunks)]

    batch_size = 100
    total_inserted = 0
    for i in range(0, len(rows), batch_size):
        batch = rows[i:i + batch_size]
        try:
            supabase.table("documents").insert(batch).execute()
            total_inserted += len(batch)
        except Exception as e:
            print(f"   ⚠️ Batch error: {e}")
            for row in batch:
                try:
                    supabase.table("documents").insert(row).execute()
                    total_inserted += 1
                except Exception as e2:
                    print(f"   ⚠️ Row error: {e2}")

    print(f"✅ Added: {filename} ({total_inserted}/{len(rows)} chunks)")
    return total_inserted > 0

def add_single_file_incremental(file_path: str, content: str, category: str = "general",
                                reindex: bool = False) -> bool:
    return add_document_incremental(file_path, content, category, Path(file_path).name, reindex=reindex)

def add_documents_from_zip_incremental(zip_path: str, reindex_existing: bool = False) -> Dict[str, int]:
    print("\n" + "="*80 + "\nINCREMENTAL DOCUMENT ADDITION\n" + "="*80)
    if not test_connection():
        print("❌ Cannot connect to Supabase.")
        return {"added": 0, "skipped": 0, "reindexed": 0, "failed": 0}

    file_paths_to_process = []
    skipped_paths = []

    with zipfile.ZipFile(zip_path, "r") as zip_ref:
        for file_path in zip_ref.namelist():
            if file_path.endswith("/") or "__MACOSX" in file_path or Path(file_path).name.startswith("._"):
                continue
            fname = Path(file_path).name.lower()
            if fname in [".ds_store", "thumbs.db", ".gitignore", ".gitkeep"]:
                continue
            if Path(file_path).suffix.lower() not in [".txt", ".pdf", ".docx"]:
                continue
            if len(Path(file_path).parts) < 3:
                continue
            if document_exists(file_path) and not reindex_existing:
                skipped_paths.append(file_path)
            else:
                file_paths_to_process.append(file_path)

    print(f"   New: {len(file_paths_to_process)} | Skipping: {len(skipped_paths)}")
    if not file_paths_to_process:
        print("✅ All documents already exist.")
        return {"added": 0, "skipped": len(skipped_paths), "reindexed": 0, "failed": 0}

    docs_by_cat = load_documents_from_zip(zip_path, file_paths_filter=set(file_paths_to_process))
    stats = {"added": 0, "skipped": len(skipped_paths), "reindexed": 0, "failed": 0}

    for cat, docs in docs_by_cat.items():
        print(f"\n📁 {cat} ({len(docs)} documents)")
        for doc in docs:
            exists = document_exists(doc["file_path"])
            if exists and reindex_existing:
                stats["reindexed"] += 1
            success = add_document_incremental(
                file_path=doc["file_path"], content=doc["content"],
                category=doc["category"], filename=doc["filename"],
                reindex=(exists and reindex_existing)
            )
            if success:
                stats["added"] += 1
            else:
                stats["failed"] += 1

    print(f"\n✅ Added: {stats['added']} | ⏭️ Skipped: {stats['skipped']} | "
          f"🔄 Re-indexed: {stats['reindexed']} | ❌ Failed: {stats['failed']}")
    return stats


# ============ Build Pipeline ============

def build_alkhidmat_rag(zip_path: str, clear_existing: bool = False, incremental: bool = False):
    print("\n" + "="*80 + "\nBUILDING ALKHIDMAT RAG SYSTEM\n" + "="*80)
    if not test_connection():
        print("❌ Cannot connect to Supabase.")
        return
    if incremental:
        add_documents_from_zip_incremental(zip_path, reindex_existing=False)
        DomainClassifier.initialize_domain_embeddings()
        print("✅ INCREMENTAL BUILD COMPLETE!")
        return
    if clear_existing:
        clear_documents_table()
    docs, meta = prepare_documents(zip_path)
    if not docs:
        print("❌ No documents found.")
        return
    chunks, chunk_meta = split_documents(docs, meta)
    embeddings = create_embeddings(chunks)
    save_chunks_to_supabase(chunks, chunk_meta, embeddings)
    DomainClassifier.initialize_domain_embeddings()
    print("✅ BUILD COMPLETE!")


# ============ Retrieval ============

def retrieve_from_supabase(query: str, top_k: int = 8, filter_category: str = None,
                           query_embedding: Optional[np.ndarray] = None) -> Tuple[List[Dict], np.ndarray, List[np.ndarray]]:
    embed_start = time.time()
    if query_embedding is not None:
        print(f"[RETRIEVAL] Using provided embedding", flush=True)
    else:
        cached_emb = find_similar_cached_embedding(query)
        if cached_emb is not None:
            query_embedding = cached_emb
            print(f"[RETRIEVAL] Using cached embedding", flush=True)
        else:
            query_embedding = get_embedder().encode([f"query: {query}"], normalize_embeddings=True)[0]
            cache_query_embedding(query, query_embedding)
    print(f"[TIMING] Query embedding: {time.time() - embed_start:.2f}s", flush=True)

    supabase = get_supabase_client()
    try:
        rpc_start = time.time()
        params = {'query_embedding': query_embedding.tolist(), 'match_threshold': RELEVANCE_THRESHOLD, 'match_count': top_k}
        if filter_category:
            params['filter_category'] = filter_category
        result = supabase.rpc('match_documents', params).execute()
        rows = result.data
        print(f"[TIMING] Supabase RPC: {time.time() - rpc_start:.2f}s", flush=True)

        if rows:
            fetch_start = time.time()
            doc_ids = [row['doc_id'] for row in rows]
            full_docs = supabase.table("documents").select("doc_id, embedding").in_("doc_id", doc_ids).execute()
            doc_map = {d['doc_id']: d for d in full_docs.data}
            for row in rows:
                if row['doc_id'] in doc_map:
                    row['embedding'] = doc_map[row['doc_id']]['embedding']
            print(f"[TIMING] Fetch embeddings: {time.time() - fetch_start:.2f}s", flush=True)

    except Exception as e:
        print(f"⚠️ RPC fallback: {e}")
        result = supabase.table("documents").select(
            "doc_id, chunk_text, category, filename, file_path, chunk_index, embedding"
        ).execute()
        rows = []
        for row in result.data:
            if row['embedding']:
                sim = float(np.dot(query_embedding, np.array(row['embedding'])))
                if sim > RELEVANCE_THRESHOLD:
                    row['similarity'] = sim
                    rows.append(row)
        rows = sorted(rows, key=lambda x: x['similarity'], reverse=True)[:top_k]

    results = []
    doc_embeddings = []
    for row in rows:
        results.append({
            "text": row['chunk_text'], "category": row['category'], "filename": row['filename'],
            "file_path": row['file_path'], "chunk_index": row.get('chunk_index', 0),
            "similarity": float(row.get('similarity', 0))
        })
        if 'embedding' in row and row['embedding'] is not None:
            try:
                emb = row['embedding']
                if isinstance(emb, str):
                    import ast
                    emb = ast.literal_eval(emb)
                doc_embeddings.append(np.array(emb, dtype=np.float32))
            except Exception as e:
                print(f"⚠️ Embedding parse error: {e}")

    print(f"\n{'='*80}", flush=True)
    print(f"RETRIEVAL: {len(results)} chunks (threshold: {RELEVANCE_THRESHOLD})", flush=True)
    for i, r in enumerate(results, 1):
        print(f" {i}. [{r['category']}] {r['filename']} (sim: {r['similarity']:.3f})", flush=True)
    print("="*80 + "\n", flush=True)
    return results, query_embedding, doc_embeddings


def sanitize_chunk_text(text: str) -> str:
    text = re.sub(r'(?mi)^\s*(user\s+question|question|q:)\s*[:\-–]?\s*.*$', '', text)
    text = re.sub(r'(?mi)^\s*(answer|a:)\s*[:\-–]?\s*.*$', '', text)
    text = re.sub(r'(?is)(?:^|\n)\s*q[:\.\-\)]\s*.*?\n\s*a[:\.\-\)]\s*.*?(?:\n|$)', '', text)
    text = re.sub(r'\[insert .*?\]', '', text, flags=re.IGNORECASE)
    text = re.sub(r'click here', '', text, flags=re.IGNORECASE)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def clean_llm_response(text: str) -> str:
    if not text:
        return text
    text = text.replace('\r\n', '\n')
    text = re.sub(r'\[?Context \d+\]?', '', text, flags=re.IGNORECASE)
    text = re.sub(r'based on Context \d+', '', text, flags=re.IGNORECASE)
    text = re.sub(r'as per (?:the )?context', '', text, flags=re.IGNORECASE)

    cutoff_patterns = [r'\nUser question\s*:', r'\nQuestion\s*:', r'\nUser question', r'\nQuestion', r'\nAnswer\s*:']
    earliest = len(text)
    for pat in cutoff_patterns:
        m = re.search(pat, text, flags=re.IGNORECASE)
        if m:
            earliest = min(earliest, m.start())
    if earliest < len(text):
        text = text[:earliest].strip()

    text = re.sub(r'(?mi)^(user question|question|answer)\s*[:\-–]\s*', '', text)

    lines = text.split('\n')
    seen = set()
    out_lines = []
    for line in lines:
        s = line.strip()
        if not s:
            out_lines.append('')
            continue
        if s.lower() in seen:
            continue
        seen.add(s.lower())
        out_lines.append(line)
    return re.sub(r'\n{3,}', '\n\n', '\n'.join(out_lines)).strip()


# ============ Answer Generation (non-selfrag path) ============

def generate_answer(query: str, top_k: int = 8, max_tokens: int = 400, filter_category: str = None):
    """Standard RAG answer generation using GPT."""
    start_time = time.time()
    print(f"[RAG] Processing: {query[:50]}...", flush=True)

    profile = build_query_lang_profile(query)
    query_info = analyze_query(profile.original_query)
    query_for_rag = profile.query_en
    original_query = profile.original_query

    if profile.input_lang == "roman_ur":
        print(f"[RAG] Roman Urdu reformulated: '{query_for_rag}'", flush=True)

    domain_classification = DomainClassifier.classify_domain(query_for_rag)
    winning_domain = domain_classification.get("domain", "general")
    retrieval_query_en = expand_query_for_retrieval(query_for_rag, winning_domain, query_info)

    results, query_embedding, doc_embeddings = retrieve_from_supabase(
        retrieval_query_en, top_k=top_k, filter_category=filter_category
    )

    if not results:
        no_answer = ("معاف کیجیے، میں اس سوال کا جواب دینے کے لیے متعلقہ معلومات نہیں ڈھونڈ سکا۔"
                     if profile.output_lang == "ur"
                     else "I apologize, but I couldn't find relevant information to answer this question.")
        return no_answer, original_query, profile.input_lang, [], {}, domain_classification

    context_parts = []
    for r in results:
        chunk_text = sanitize_chunk_text(r["text"])
        if profile.output_lang in ["ur", "roman_ur"] and TRANSLATE_CONTEXT_FOR_URDU_OUTPUT:
            try:
                chunk_text = translate_english_to_urdu(chunk_text)
            except Exception:
                pass
        if chunk_text:
            if len(chunk_text) > 1200:
                chunk_text = chunk_text[:1200].rsplit('\n', 1)[0] + "\n\n[truncated]"
            context_parts.append(chunk_text)
    context = "\n\n".join(context_parts)

    wants_list = query_info['wants_list']
    wants_summary = query_info['wants_summary']
    wants_detail = query_info['wants_detail']

    if profile.output_lang == "ur":
        format_instruction = ("براہ کرم نقاط کی شکل میں واضح جواب دیں۔" if wants_list
                              else "براہ کرم 2-3 جملوں میں مختصر خلاصہ دیں۔" if wants_summary
                              else "براہ کرم مکمل اور واضح جواب دیں۔")
        display_question = profile.query_urdu_script or original_query
        history_block_ur = ""
        if history:
            lines = []
            for turn in history[-6:]:
                role = "صارف" if turn.get("role") == "user" else "اسسٹنٹ"
                lines.append(f"{role}: {(turn.get('content') or '').strip()}")
            history_block_ur = "\nگفتگو:\n" + "\n".join(lines) + "\n"
 
        prompt = f"""آپ الخدمت فاؤنڈیشن پاکستان کے لیے ایک مددگار اسسٹنٹ ہیں۔
{format_instruction}
{history_block_ur}
ہدایات:
- صرف نیچے دیے گئے سیاق و سباق کی معلومات استعمال کریں۔
- اگر جواب سیاق و سباق میں نہیں تو صرف لکھیں: "مجھے معلوم نہیں"
- جواب صرف اردو (نستعلیق) میں دیں — انگریزی یا رومن اردو ہرگز نہ لکھیں۔
 
سیاق و سباق:
{context}
 
سوال: {display_question}
 
جواب (صرف اردو میں):
"""
    else:
        format_instruction = ("Provide a clear answer in bullet point format." if wants_list
                              else "Provide a brief summary in 2-3 sentences." if wants_summary
                              else "Provide a detailed, comprehensive answer." if wants_detail
                              else "Provide a clear, complete answer.")
        prompt = f"""You are a helpful customer support agent for Alkhidmat Foundation Pakistan.

{format_instruction}

CRITICAL INSTRUCTIONS:
- Use ONLY the information in the context below
- DO NOT answer from general knowledge
- If the answer is not in the context, respond EXACTLY: "I don't know"
- Return ONLY the direct answer, no labels

Information:
{context}

User question: {original_query}

Answer:
"""

    print(f"[RAG] Generating answer (GPT: {GPT_MODEL})...", flush=True)
    answer = _gpt_generate(prompt, max_tokens=max_tokens)
    answer = clean_llm_response(answer)

    retrieval_conf = ConfidenceScorer.calculate_retrieval_confidence(query_embedding, doc_embeddings, top_k)
    confidence_scores = ConfidenceScorer.calculate_combined_confidence(retrieval_confidence=retrieval_conf)

    # Print confidence scores right after generation
    _print_confidence_scores(confidence_scores)

    if profile.output_lang == "ur":
        if not is_urdu_script(answer):
            answer = translate_english_to_urdu(answer, timeout=15)
    elif profile.output_lang == "roman_ur":
        protected = protect_brand_terms(answer)
        answer_ur = translate_english_to_urdu(protected, timeout=15)
        answer = to_roman_urdu(restore_brand_terms(answer_ur))

    sources = [{"category": r['category'], "filename": r['filename'],
                "file_path": r['file_path'], "similarity": r['similarity']} for r in results]

    print(f"[TIMING] Total: {time.time() - start_time:.2f}s", flush=True)
    return answer, original_query, profile.input_lang, sources, confidence_scores, domain_classification



# ── Failure routing helper ─────────────────────────────────────────────────
_AGENT_ROUTE_SIGNAL = "__ROUTE_TO_AGENT__"

def _agent_route_response(profile) -> str:
    """
    Return a language-appropriate "routing to agent" message.
    The API detects __ROUTE_TO_AGENT__ sentinel in input_lang and creates a ticket.
    We embed the sentinel in the returned tuple via selfrag_metrics['route_to_agent'].
    """
    if profile.output_lang == "ur":
        return "معاف کیجیے، میں آپ کے سوال کا جواب نہیں دے سکا۔ میں آپ کو ایک انسانی ایجنٹ سے منسلک کر رہا ہوں۔"
    elif profile.output_lang == "roman_ur":
        return "Maafi chahta/chahti hun, main aapke sawal ka jawab nahi de saka/saki. Main aapko ek insani agent se connect kar raha/rahi hun."
    else:
        return "I wasn't able to answer your question. Routing you to a human agent."

# ============ CONVERSATION MEMORY — QUERY REWRITING ============
 
def rewrite_query_with_history(
    current_query_en: str,
    conversation_history: list,
) -> str:
    """
    Rewrites the current (English) retrieval query into a self-contained
    question using the last few conversation turns.
 
    Implements the LCEL conversational RAG pattern (Harrison Chase, 2024):
    history-conditioned query rewriting keeps retrieval stable across
    multi-turn conversations without stuffing raw history into the
    vector search.
 
    Args:
        current_query_en:    The English form of the current query
                             (already translated/transliterated).
        conversation_history: [{"role": "user"|"assistant", "content": str}, ...]
                             Ordered oldest-first, most-recent last.
 
    Returns:
        Rewritten English query, or current_query_en unchanged on failure.
    """
    if not conversation_history:
        return current_query_en
 
    recent = conversation_history[-6:]   # last 3 user + 3 assistant turns
 
    history_text = ""
    for turn in recent:
        role = turn.get("role", "user")
        content = (turn.get("content") or "").strip()
        label = "User" if role == "user" else "Assistant"
        history_text += f"{label}: {content}\n"
 
    prompt = f"""Given the conversation history and a follow-up question, rewrite the follow-up question into a fully self-contained English question that can be understood without the history.
 
Rules:
- Resolve all pronouns and coreferences (e.g. "it", "that program", "there").
- Output ONLY the rewritten question — no preamble or explanation.
- Keep it under 25 words if possible.
- If already self-contained, return it unchanged.
- Always output English regardless of the follow-up language.
 
Conversation history:
{history_text.strip()}
 
Follow-up question: {current_query_en}
 
Rewritten standalone question:"""
 
    rewritten = _gpt_call(prompt, max_tokens=60, temperature=0.1)
    rewritten = (rewritten or "").strip().strip("\"'")
 
    if rewritten:
        print(f"[MEMORY] Rewritten for retrieval: '{current_query_en}' → '{rewritten}'", flush=True)
        return rewritten
 
    return current_query_en

# ============ SELF-RAG ANSWER GENERATION ============

def generate_answer_selfrag(query: str, top_k: int = 8, max_tokens: int = 400,
                              filter_category: str = None,
                              conversation_history: list = None):
    """
    Multilingual Self-RAG with Agentic AI.
    ALL LLM calls use GPT via _gpt_call / _gpt_generate.

    Pipeline:
      Step 0.5 — Router: greetings, niceties, retrieval necessity (GPT)
      Step 1   — Domain relevance check (GPT)
      Step 2   — Retrieve + embedding cache + agentic retry
      Step 3   — Similarity filter (pure embedding, NO GPT call)
      Step 3.5 — Pre-generation gate: can context answer query? (GPT)
      Step 4   — Answer generation (GPT)
      [print confidence scores immediately after generation]
      Step 5   — Post-generation hallucination check: verify_support (GPT)
      Step 4.5 — Evidence coverage: claim-by-claim (GPT), skipped if fully_supported
      Step 6   — Utility evaluation (GPT)
    """
    start_time = time.time()
    print(f"\n{'='*80}", flush=True)
    print(f"AGENTIC SELF-RAG  [GPT: {GPT_MODEL}]", flush=True)
    print(f"{'='*80}", flush=True)

    profile = build_query_lang_profile(query)

    # ── Conversation memory: rewrite retrieval query ──────────────────
    history = conversation_history or []
    if history:
        rewritten_en = rewrite_query_with_history(profile.query_en, history)
        profile.query_en = rewritten_en   # direct field assignment — no dataclasses.replace needed
    # ─────────────────────────────────────────────────────────────────

    query_info = analyze_query(profile.original_query)
    query_for_rag = profile.query_en
    original_query = profile.original_query

    if profile.input_lang == "roman_ur":
        print(f"[AGENTIC-RAG] Roman Urdu GPT-reformulated: '{query_for_rag}'", flush=True)

    critic = SelfRAGCritic()
    router_agent = RouterAgent(critic)
    retriever_agent = RetrieverAgent()
    evidence_agent = EvidenceCoverageAgent(critic)
    selfrag_metrics = {
        'domain_relevant': True, 'retrieve_needed': False,
        'answer_in_context': False, 'relevance_score': 0.0,
        'support_level': 'uncertain', 'utility_rating': 0,
        'embedding_cached': False, 'retrieval_retried': False,
        'evidence_coverage': 1.0,
        'roman_ur_reformulated': profile.input_lang == "roman_ur",
        'perspective_reformulated': False,
        'route_to_agent': False
    }

    # STEP 0.5: Router (GPT)
    print(f"\n[AGENTIC-RAG] Router Agent (GPT)...", flush=True)
    should_retrieve, router_confidence, cached_answer = router_agent.route(query_for_rag)
    selfrag_metrics['retrieve_needed'] = should_retrieve

    if not should_retrieve and cached_answer:
        print(f"[AGENTIC-RAG] ✅ Router: no retrieval needed", flush=True)
        return (cached_answer, original_query, profile.input_lang, [],
                {'combined_confidence': router_confidence}, {}, selfrag_metrics)

    # STEP 1: Domain relevance (GPT)
    if SELFRAG_ENABLE:
        print(f"\n[SELF-RAG] Step 1: Domain relevance check (GPT)...", flush=True)
        is_domain_relevant = critic.is_domain_relevant(query_for_rag)
        selfrag_metrics['domain_relevant'] = is_domain_relevant
        if not is_domain_relevant:
            print(f" ✗ IRRELEVANT to Alkhidmat Foundation", flush=True)
            dummy_cls = {'domain': 'irrelevant', 'confidence': 0.0,
                         'all_scores': {'donation': 0.0, 'healthcare': 0.0, 'general': 0.0}}
            return ("That is an irrelevant question.", original_query, profile.input_lang,
                    [], {'combined_confidence': 0.0}, dummy_cls, selfrag_metrics)
        print(f" ✓ RELEVANT", flush=True)

    selfrag_metrics['retrieve_needed'] = True

    # STEP 1.5: Reformulate query from Alkhidmat's perspective (improves retrieval
    # for queries that omit the org name, e.g. "what are your services?")
    print(f"\n[SELF-RAG] Step 1.5: Alkhidmat-perspective reformulation...", flush=True)
    query_for_rag = retriever_agent.reformulate_for_alkhidmat_perspective(query_for_rag)
    selfrag_metrics['perspective_reformulated'] = True

    # Domain classification (embedding-based, no GPT)
    domain_classification = DomainClassifier.classify_domain(query_for_rag)
    winning_domain = domain_classification.get("domain", "general")
    print(f"\n[SELF-RAG] Domain: {winning_domain} ({domain_classification['confidence']:.2%})", flush=True)

    retrieval_query_en = expand_query_for_retrieval(query_for_rag, winning_domain, query_info)

    # STEP 2: Embedding cache + retrieval
    # Skip cache when conversation memory rewrote the query — the rewritten query
    # targets a different topic than the previous turn and must not reuse its embedding.
    _memory_rewrote_query = bool(history) and (retrieval_query_en != profile.original_query)
    print(f"\n[AGENTIC-RAG] Embedding cache check...", flush=True)
    cached_embedding = None if _memory_rewrote_query else find_similar_cached_embedding(retrieval_query_en)
    if cached_embedding is not None:
        query_embedding = cached_embedding
        selfrag_metrics['embedding_cached'] = True
        print(f"[AGENTIC-RAG] ✅ Cache hit!", flush=True)
    else:
        if _memory_rewrote_query:
            print(f"[AGENTIC-RAG] Cache bypassed — memory-rewritten query needs fresh embedding", flush=True)
        query_embedding = get_embedder().encode([f"query: {retrieval_query_en}"], normalize_embeddings=True)[0]
        cache_query_embedding(retrieval_query_en, query_embedding)

    print(f"\n[SELF-RAG] Step 2: Retrieving documents...", flush=True)
    results, query_embedding, doc_embeddings, was_retried = retriever_agent.retrieve_with_retry(
        retrieval_query_en, winning_domain, top_k=top_k, filter_category=filter_category
    )
    selfrag_metrics['retrieval_retried'] = was_retried

    if not results:
        print(f"[SELF-RAG] ⚠️ No documents found — routing to agent", flush=True)
        selfrag_metrics['route_to_agent'] = True
        return (_agent_route_response(profile), original_query, profile.input_lang,
                [], {'combined_confidence': 0.0, 'route_to_agent': True}, domain_classification, selfrag_metrics)

    # STEP 3: Similarity filter — pure embedding, NO GPT call
    # assess_relevance() has been REMOVED. Supabase pre-filters by RELEVANCE_THRESHOLD.
    # A tighter threshold here drops only the weakest docs with zero extra GPT cost.
    STEP3_SIMILARITY_THRESHOLD = float(os.environ.get("STEP3_SIMILARITY_THRESHOLD", "0.72"))
    if SELFRAG_ENABLE:
        print(f"\n[SELF-RAG] Step 3: Similarity filter (threshold: {STEP3_SIMILARITY_THRESHOLD}, NO GPT)...", flush=True)
        relevant_results = []
        relevance_scores = []
        for i, result in enumerate(results):
            sim = result.get('similarity', 0.0)
            relevance_scores.append(sim)
            if sim >= STEP3_SIMILARITY_THRESHOLD:
                relevant_results.append(result)
                print(f" ✓ Doc {i+1}: kept   (sim: {sim:.3f})", flush=True)
            else:
                print(f" ✗ Doc {i+1}: dropped (sim: {sim:.3f} < {STEP3_SIMILARITY_THRESHOLD})", flush=True)

        if not relevant_results:
            print(f" ⚠️ All dropped — keeping all {len(results)} (all passed base threshold)", flush=True)
            relevant_results = results

        results = relevant_results
        avg_relevance = float(np.mean(relevance_scores)) if relevance_scores else 0.0
        selfrag_metrics['relevance_score'] = avg_relevance
        print(f" → Kept {len(results)} docs | avg sim: {avg_relevance:.3f}", flush=True)

    # Build context
    context_parts = []
    for r in results:
        chunk_text = sanitize_chunk_text(r["text"])
        if not chunk_text:
            continue
        if profile.output_lang == "ur" and TRANSLATE_CONTEXT_FOR_URDU_OUTPUT:
            try:
                chunk_text = translate_english_to_urdu(chunk_text)
            except Exception:
                pass
        if len(chunk_text) > 1200:
            chunk_text = chunk_text[:1200].rsplit('\n', 1)[0] + "\n\n[truncated]"
        context_parts.append(chunk_text)
    context = "\n\n".join(context_parts)

    # STEP 3.5: Pre-generation gate (GPT)
    if SELFRAG_ENABLE:
        print(f"\n[SELF-RAG] Step 3.5: Pre-generation gate (GPT)...", flush=True)
        top_sim = results[0].get('similarity', 0) if results else 0
        if top_sim >= 0.86:
            print(f" ✓ Skipping gate — top doc similarity is high ({top_sim:.3f})", flush=True)
            can_answer = True
            selfrag_metrics['answer_in_context'] = True
        else:
            can_answer = critic.check_answer_in_context(query_for_rag, context)
            selfrag_metrics['answer_in_context'] = can_answer
        if not can_answer:
            # Context mismatch — likely a cache/retrieval topic drift. Retry once with
            # a fresh, domain-specific reformulation before routing to a human agent.
            print(f" ✗ Context cannot answer — retrying retrieval with fresh query...", flush=True)
            retry_query = retriever_agent.reformulate_query(
                query_for_rag, winning_domain, reason="context_mismatch"
            )
            retry_embedding = get_embedder().encode([f"query: {retry_query}"], normalize_embeddings=True)[0]
            cache_query_embedding(retry_query, retry_embedding)
            retry_results, retry_embedding, retry_doc_embeddings, _ = retriever_agent.retrieve_with_retry(
                retry_query, winning_domain, top_k=top_k, filter_category=filter_category
            )
            if retry_results:
                retry_context_parts = []
                for r in retry_results:
                    chunk = sanitize_chunk_text(r["text"])
                    if chunk:
                        if len(chunk) > 1200:
                            chunk = chunk[:1200].rsplit('\n', 1)[0] + "\n\n[truncated]"
                        retry_context_parts.append(chunk)
                retry_context = "\n\n".join(retry_context_parts)
                can_answer_retry = critic.check_answer_in_context(query_for_rag, retry_context)
                if can_answer_retry:
                    print(f" ✅ Retry retrieval succeeded — continuing with new context", flush=True)
                    results = retry_results
                    query_embedding = retry_embedding
                    doc_embeddings = retry_doc_embeddings
                    context = retry_context
                    selfrag_metrics['retrieval_retried'] = True
                    selfrag_metrics['answer_in_context'] = True
                else:
                    print(f" ✗ Retry also failed — routing to agent", flush=True)
                    selfrag_metrics['route_to_agent'] = True
                    return (_agent_route_response(profile), original_query, profile.input_lang,
                            results, {'combined_confidence': 0.0, 'route_to_agent': True}, domain_classification, selfrag_metrics)
            else:
                print(f" ✗ Retry retrieval returned no results — routing to agent", flush=True)
                selfrag_metrics['route_to_agent'] = True
                return (_agent_route_response(profile), original_query, profile.input_lang,
                        results, {'combined_confidence': 0.0, 'route_to_agent': True}, domain_classification, selfrag_metrics)
        else:
            print(f" ✓ Context can answer the query", flush=True)

    # Build prompt
    wants_list = query_info['wants_list']
    wants_summary = query_info['wants_summary']
    wants_detail = query_info['wants_detail']

    if profile.output_lang == "ur":
        format_instruction = ("براہ کرم نقاط کی شکل میں واضح جواب دیں۔" if wants_list
                              else "براہ کرم 2-3 جملوں میں مختصر خلاصہ دیں۔" if wants_summary
                              else "براہ کرم مکمل اور واضح جواب دیں۔")
        display_question = (profile.query_urdu_script
                            if profile.input_lang == "roman_ur" and profile.query_urdu_script
                            else original_query)
        prompt = f"""آپ الخدمت فاؤنڈیشن پاکستان کے لیے ایک مددگار اسسٹنٹ ہیں۔
{format_instruction}
صرف نیچے دیے گئے سیاق و سباق کی معلومات استعمال کریں۔ اگر جواب سیاق و سباق میں نہیں تو کہیں: "مجھے معلوم نہیں"

سیاق و سباق:
{context}

سوال: {display_question}

جواب (صرف اردو میں):
"""
    else:
        # Inject recent conversation turns so the model can reference prior answers
        history_block = ""
        if history:
            lines = []
            for turn in history[-6:]:
                role = "User" if turn.get("role") == "user" else "Assistant"
                lines.append(f"{role}: {(turn.get('content') or '').strip()}")
            history_block = "\nConversation so far:\n" + "\n".join(lines) + "\n"

        if profile.output_lang == "roman_ur":
            # Dedicated Roman Urdu prompt — language instruction is the FIRST thing GPT sees,
            # not a footnote. This dramatically reduces English bleed-through.
            format_instruction_ru = (
                "Jawab nuke (bullet points) mein dein." if wants_list
                else "2-3 jumlon mein mukhtasar jawab dein." if wants_summary
                else "Mukammal aur tafsili jawab dein." if wants_detail
                else "Saaf aur mukammal jawab dein."
            )
            prompt = f"""Aap Alkhidmat Foundation Pakistan ke liye ek madad-gar assistant hain.
ZAROORI HUKAM: Sirf Roman Urdu mein jawab dein (Latin haroof mein — koi Urdu/Arabic script nahi).
Misaal: "Alkhidmat Foundation mein aap online donate kar sakte hain..."
{format_instruction_ru}
{history_block}
AHKAM:
- Sirf neeche diye gaye malumat se jawab dein.
- Apni taraf se kuch mat bataein.
- Agar malumat mojood nahi: "Mujhe maloom nahin" likhein.
- Sirf jawab likhein, koi label nahi.

Malumat:
{context}

Sawal: {original_query}

Jawab (Roman Urdu mein):"""
        else:
            format_instruction = (
                "Provide a clear answer in bullet point format." if wants_list
                else "Provide a brief summary in 2-3 sentences." if wants_summary
                else "Provide a detailed, comprehensive answer." if wants_detail
                else "Provide a clear, complete answer."
            )
            prompt = f"""You are a helpful customer support agent for Alkhidmat Foundation Pakistan.

{format_instruction}
{history_block}
CRITICAL INSTRUCTIONS:
- Use ONLY the information provided in the context below.
- DO NOT answer from general knowledge.
- If the answer is not in the context, respond EXACTLY: "I don't know"
- Return ONLY the direct answer, no labels.
LANGUAGE RULE: Reply in English only.

Information:
{context}

User question: {original_query}

Answer:
"""

    # STEP 4: Generate answer (GPT)
    llm_start = time.time()
    print(f"\n[SELF-RAG] Step 4: Generating answer (GPT: {GPT_MODEL})...", flush=True)
    answer = _gpt_generate(prompt, max_tokens=max_tokens)
    print(f"[TIMING] GPT generation: {time.time() - llm_start:.2f}s", flush=True)
    answer = clean_llm_response(answer)

    # PRINT CONFIDENCE SCORES IMMEDIATELY after generation
    retrieval_conf_early = ConfidenceScorer.calculate_retrieval_confidence(query_embedding, doc_embeddings, top_k)
    confidence_scores_early = ConfidenceScorer.calculate_combined_confidence(retrieval_confidence=retrieval_conf_early)
    _print_confidence_scores(confidence_scores_early, label="CONFIDENCE SCORES (post-generation)")

    # Translate to English for critic steps if answer is Urdu/Roman-Urdu
    answer_for_critic = answer
    if profile.output_lang == "ur":
        try:
            answer_for_critic = translate_urdu_to_english(answer)
        except Exception:
            answer_for_critic = answer
    elif profile.output_lang == "roman_ur":
        try:
            answer_for_critic = translate_auto_to_english(answer)
        except Exception:
            answer_for_critic = answer

    # STEP 5: Post-generation hallucination check (GPT)
    # WHY BOTH Step 3.5 AND Step 5:
    # - Step 3.5 (check_answer_in_context): BEFORE generation — "is the answer findable?"
    # - Step 5 (verify_support): AFTER generation — "did the LLM stay grounded?"
    # They catch different failure modes. Both are intentionally kept.
    
    # STEP 5: Post-generation hallucination check (GPT)
    if SELFRAG_ENABLE:
        print(f"\n[SELF-RAG] Step 5: Post-generation support check (GPT)...", flush=True)
        support_level = critic.verify_support(query_for_rag, answer_for_critic, context)
        selfrag_metrics['support_level'] = support_level
        print(f" → Support: {support_level.upper()}", flush=True)
        
        # Only route to agent if NO_SUPPORT *and* retrieval confidence is also low
        # High retrieval confidence means we found the right docs — trust the answer
        if support_level == "no_support" and retrieval_conf_early < 0.82:
            print(f" ⚠️ Answer NOT supported and low retrieval confidence — routing to agent", flush=True)
            selfrag_metrics['route_to_agent'] = True
            return (_agent_route_response(profile), original_query, profile.input_lang,
                    results, {'combined_confidence': 0.0, 'route_to_agent': True}, domain_classification, selfrag_metrics)
        elif support_level == "no_support":
            print(f" ⚠️ NO_SUPPORT but retrieval confidence is high ({retrieval_conf_early:.3f}) — proceeding anyway", flush=True)

    # STEP 4.5: Evidence coverage (GPT), skipped if fully supported
    if EVIDENCE_COVERAGE_ENABLE:
        if selfrag_metrics.get('support_level') == 'fully_supported':
            print(f"\n[AGENTIC-RAG] Evidence Coverage: skipping — already fully supported ✅", flush=True)
            selfrag_metrics['evidence_coverage'] = 1.0
        else:
            print(f"\n[AGENTIC-RAG] Evidence Coverage: claim-by-claim check (GPT)...", flush=True)
            all_covered, unsupported_claims, coverage_score = evidence_agent.check_coverage(
                answer, context, query_for_rag, domain=winning_domain, results=results
            )
            selfrag_metrics['evidence_coverage'] = coverage_score
            allow_summary = evidence_agent._is_summary_allowed(query_for_rag, winning_domain)
            if not all_covered and coverage_score < (0.3 if allow_summary else 0.4):
                print(f"[AGENTIC-RAG] ⚠️ Low coverage ({coverage_score:.2f}) — removing unsupported claims...", flush=True)
                sentences = re.split(r'[.!?]\s+', answer)
                filtered = [s for s in sentences
                            if not any(claim.lower() in s.lower() for claim in unsupported_claims)]
                answer = '. '.join(filtered).strip() if filtered else ""
                if answer and not answer.endswith(('.', '!', '?')):
                    answer += '.'
                answer = re.sub(r'\s+', ' ', answer).strip()
                if not answer:
                    answer = "I cannot provide a reliable answer as some claims cannot be verified."

    # STEP 6: Utility evaluation (GPT)
    print(f"\n[SELF-RAG] Step 6: Utility evaluation (GPT)...", flush=True)
    utility_rating = critic.evaluate_utility(query_for_rag, answer_for_critic)
    selfrag_metrics["utility_rating"] = utility_rating
    print(f" → Utility: {utility_rating}/5", flush=True)

    min_utility = 2 if profile.output_lang in ("ur", "roman_ur") else 3
    if utility_rating < min_utility:
        strong_evidence = (
            bool(selfrag_metrics.get("answer_in_context")) and
            selfrag_metrics.get("support_level") in ("fully_supported", "partially_supported") and
            float(selfrag_metrics.get("evidence_coverage", 0.0) or 0.0) >= 0.80
        )
        if strong_evidence:
            print(" ⚠️ Utility low but evidence strong — NOT rejecting", flush=True)
        else:
            print(" ⚠️ Utility too low — routing to agent", flush=True)
            selfrag_metrics['route_to_agent'] = True
            return (_agent_route_response(profile), original_query, profile.input_lang,
                    [], {'combined_confidence': 0.0, 'route_to_agent': True}, domain_classification, selfrag_metrics)

    # Final confidence scores
    retrieval_conf = ConfidenceScorer.calculate_retrieval_confidence(query_embedding, doc_embeddings, top_k)
    selfrag_scores = {'relevance_score': selfrag_metrics.get('relevance_score', 0.0)}
    confidence_scores = ConfidenceScorer.calculate_combined_confidence(
        retrieval_confidence=retrieval_conf, selfrag_scores=selfrag_scores
    )

    combined_conf = confidence_scores.get('combined_confidence', 0)
    if SELFRAG_ENABLE and combined_conf < SELFRAG_MIN_CONFIDENCE:
        print(f"\n⚠️ Confidence ({combined_conf:.2f}) < threshold ({SELFRAG_MIN_CONFIDENCE}) — routing to agent", flush=True)
        selfrag_metrics['route_to_agent'] = True
        confidence_scores['route_to_agent'] = True
        return (_agent_route_response(profile), original_query, profile.input_lang,
                results, confidence_scores, domain_classification, selfrag_metrics)

    # Output language post-processing
    if profile.output_lang == "ur":
        if not is_urdu_script(answer):
            answer = translate_english_to_urdu(answer, timeout=15)
    elif profile.output_lang == "roman_ur":
        if not is_urdu_script(answer):
            # Answer came back in English — convert directly to Roman Urdu via GPT.
            # The old chain (English→Urdu script→to_roman_urdu) broke silently when
            # romanize_to_roman_urdu_with_llm failed (gpt-5.2 / local LLM unavailable).
            roman_prompt = (
                "Translate the following English text into Roman Urdu "
                "(Urdu written in Latin/English letters).\n\n"
                "STRICT RULES:\n"
                "- Output MUST be in Roman Urdu (Latin letters only). No Urdu/Arabic script.\n"
                "- Use natural Pakistani Roman Urdu: aap, hain, karein, nahin, hai, ka, ki, ke, mein, se.\n"
                "- Keep brand names unchanged: Alkhidmat, EasyPaisa, JazzCash, Bank of Punjab.\n"
                "- Keep phone numbers, URLs, and account numbers exactly as-is.\n"
                "- Output ONLY the Roman Urdu translation. No labels, no preamble.\n\n"
                f"English text:\n{answer}\n\nRoman Urdu:"
            )
            roman_answer = _gpt_call(roman_prompt, max_tokens=600, temperature=0.2)
            roman_answer = (roman_answer or "").strip()
            if roman_answer and not is_urdu_script(roman_answer):
                answer = roman_answer
            # else: GPT failed — keep English as readable fallback
        else:
            # Answer is already in Urdu script — use existing romanization path
            protected = protect_brand_terms(answer)
            answer = to_roman_urdu(restore_brand_terms(protected))

    sources = [{"category": r['category'], "filename": r['filename'],
                "file_path": r['file_path'], "similarity": r['similarity']} for r in results]

    total_time = time.time() - start_time
    print(f"\n[SELF-RAG] ✅ Answer accepted | {total_time:.2f}s total", flush=True)

    print(f"\n{'─'*80}", flush=True)
    print(f"SELF-RAG PIPELINE METRICS  [GPT: {GPT_MODEL}]", flush=True)
    print(f"{'─'*80}", flush=True)
    print(f" Domain relevant         : {'✅' if selfrag_metrics.get('domain_relevant') else '❌'}", flush=True)
    print(f" Roman Urdu reformulated : {'✅' if selfrag_metrics.get('roman_ur_reformulated') else '—'}", flush=True)
    print(f" Perspective reformulated: {'✅' if selfrag_metrics.get('perspective_reformulated') else '—'}", flush=True)
    print(f" Routed to agent         : {'🔄' if selfrag_metrics.get('route_to_agent') else '—'}", flush=True)
    print(f" Embedding cached        : {'✅' if selfrag_metrics.get('embedding_cached') else '—'}", flush=True)
    print(f" Retrieval retried       : {'🔄' if selfrag_metrics.get('retrieval_retried') else '—'}", flush=True)
    print(f" Avg doc similarity      : {selfrag_metrics.get('relevance_score', 0.0):.3f}", flush=True)
    print(f" Answer in context       : {'✅' if selfrag_metrics.get('answer_in_context') else '❌'}", flush=True)
    print(f" Support level           : {selfrag_metrics.get('support_level', 'n/a').upper()}", flush=True)
    print(f" Evidence coverage       : {selfrag_metrics.get('evidence_coverage', 0.0):.3f}", flush=True)
    print(f" Utility rating          : {selfrag_metrics.get('utility_rating', 0)}/5", flush=True)
    _print_confidence_scores(confidence_scores, label="FINAL CONFIDENCE SCORES")

    return answer, original_query, profile.input_lang, sources, confidence_scores, domain_classification, selfrag_metrics


# ============ CLI Functions ============

def query_alkhidmat_rag(query: str, category: str = None, use_selfrag: bool = True):
    if use_selfrag and SELFRAG_ENABLE:
        answer, original_query, input_lang, sources, confidence_scores, domain_classification, selfrag_metrics = \
            generate_answer_selfrag(query, filter_category=category)
    else:
        answer, original_query, input_lang, sources, confidence_scores, domain_classification = \
            generate_answer(query, filter_category=category)
        selfrag_metrics = {}

    print("\n" + "="*80)
    print("QUESTION:", original_query)
    print(f"(Detected input language: {input_lang})")
    if category:
        print("CATEGORY FILTER:", category)
    print("="*80)

    domain_emoji = DomainClassifier.get_domain_emoji(domain_classification['domain'])
    print(f"\n{domain_emoji} DOMAIN: {domain_classification['domain'].upper()} ({domain_classification['confidence']:.2%})")
    for domain, score in domain_classification['all_scores'].items():
        bar = '█' * int(score * 50)
        print(f" {DomainClassifier.get_domain_emoji(domain)} {domain:12s}: {score:.4f} {bar}")

    print("\nANSWER:")
    print(answer)
    print("\n" + "="*80 + "\n")
    return answer


def show_statistics():
    supabase = get_supabase_client()
    print("\n" + "="*80 + "\nDATABASE STATISTICS\n" + "="*80)
    try:
        result = supabase.table("documents").select("doc_id", count="exact").execute()
        print(f"\nTotal chunks: {result.count}")
        result = supabase.table("documents").select("category").execute()
        categories = {}
        for row in result.data:
            cat = row.get('category', 'Unknown')
            categories[cat] = categories.get(cat, 0) + 1
        print("\nDocuments by Category:")
        for cat, count in sorted(categories.items(), key=lambda x: x[1], reverse=True):
            print(f" {cat}: {count} chunks")
    except Exception as e:
        print(f"Error: {e}")


def batch_query_file(input_file: str, output_file: str, use_selfrag: bool = True):
    os.environ['BATCH_MODE'] = 'True'
    print(f"\n{'='*80}\nBATCH QUERY: {input_file} → {output_file}\n{'='*80}")
    try:
        with open(input_file, 'r', encoding='utf-8') as f:
            queries = [line.strip() for line in f if line.strip()]
        print(f"Found {len(queries)} queries.")

        with open(output_file, 'w', encoding='utf-8') as f_out:
            f_out.write(f"ALKHIDMAT RAG (SUPABASE + GPT: {GPT_MODEL}) — BATCH RESULTS\n")
            f_out.write(f"Date: {time.strftime('%Y-%m-%d %H:%M:%S')}\n\n")

            for i, query in enumerate(queries):
                print(f"Processing {i+1}/{len(queries)}: '{query[:40]}...'")
                if use_selfrag and SELFRAG_ENABLE:
                    answer, original_query, input_lang, sources, confidence_scores, domain_classification, selfrag_metrics = \
                        generate_answer_selfrag(query)
                else:
                    answer, original_query, input_lang, sources, confidence_scores, domain_classification = \
                        generate_answer(query)
                    selfrag_metrics = {}

                f_out.write(f"{'='*80}\nQUERY {i+1}/{len(queries)}\n{'='*80}\n")
                f_out.write(f"QUERY: {original_query}\nLANGUAGE: {input_lang}\n")
                f_out.write(f"DOMAIN: {domain_classification['domain'].upper()} ({domain_classification['confidence']:.2%})\n\n")
                f_out.write(f"ANSWER:\n{answer}\n\n")
                f_out.write(f"CONFIDENCE: {confidence_scores.get('combined_confidence', 0):.4f} "
                            f"(retrieval: {confidence_scores.get('retrieval_confidence', 0):.4f})\n")
                f_out.write(f"SOURCES:\n")
                if sources:
                    for s in sources:
                        f_out.write(f" - [{s['category']}] {s['filename']} (sim: {s['similarity']:.3f})\n")
                else:
                    f_out.write(" - No relevant documents found.\n")
                f_out.write(f"\n{'='*40}\n\n")
                print(f" ✓ Domain: {domain_classification['domain']} | Conf: {confidence_scores.get('combined_confidence', 0):.2f}")

        print(f"\n✅ BATCH COMPLETE → {output_file}")
    except FileNotFoundError:
        print(f"❌ File not found: {input_file}")
    except Exception as e:
        print(f"❌ Batch error: {e}")
    finally:
        os.environ['BATCH_MODE'] = 'False'


# ============ Main ============

if __name__ == "__main__":
    DEFAULT_ZIP = "Al Khidmat Knowledge Base.zip"

    if len(sys.argv) > 1 and sys.argv[1] == "test":
        test_connection()

    elif len(sys.argv) > 1 and sys.argv[1] == "build":
        zip_path = sys.argv[2] if len(sys.argv) > 2 else DEFAULT_ZIP
        build_alkhidmat_rag(zip_path, clear_existing="--clear" in sys.argv)

    elif len(sys.argv) > 1 and sys.argv[1] == "query":
        q = " ".join(sys.argv[2:]) if len(sys.argv) > 2 else "What donation methods does Alkhidmat accept?"
        query_alkhidmat_rag(q, use_selfrag="--no-selfrag" not in sys.argv)

    elif len(sys.argv) > 1 and sys.argv[1] == "file_query":
        batch_query_file(
            sys.argv[2] if len(sys.argv) > 2 else "input_queries.txt",
            sys.argv[3] if len(sys.argv) > 3 else "output_answers.txt",
            use_selfrag="--no-selfrag" not in sys.argv
        )

    elif len(sys.argv) > 1 and sys.argv[1] == "stats":
        show_statistics()

    else:
        print(f"\n{'='*80}")
        print(f"ALKHIDMAT RAG SYSTEM (SUPABASE + GPT: {GPT_MODEL})")
        print("="*80)
        print("\nUSAGE:")
        print(" python RAG_supabase_enhanced.py test")
        print(" python RAG_supabase_enhanced.py build [zip_path] [--clear]")
        print(" python RAG_supabase_enhanced.py query 'your question'")
        print(" python RAG_supabase_enhanced.py file_query input.txt output.txt")
        print(" python RAG_supabase_enhanced.py stats")
        print("\nENV VARS:")
        print(" OPENAI_API_KEY=sk-...          (required)")
        print(" GPT_MODEL=gpt-4o-mini          (default: gpt-4o-mini)")
        print(" STEP3_SIMILARITY_THRESHOLD=0.72 (default)")
        print("="*80)